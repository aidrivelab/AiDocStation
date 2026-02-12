# -*- coding: utf-8 -*-
"""
@File    : local_bridge/utils/reference_doc.py
@Desc    : AiDoc Station Lite 核心模块 - 赋能高效文档协作与智能排版处�?
@Author  : AIDriveLab Team
@Create  : 2026-02-09 21:12:43
@Version : V0.2.6
@Copyright: ©AIDriveLab Inc. All Rights Reserved.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



def _create_style_with_id(doc, name_cn: str, style_id: str, base_style: str = None):
    







       

    try:
        style = doc.styles[style_id]
        style.name = name_cn
        return style
    except KeyError:
        pass
        

    try:
        style = doc.styles[name_cn]

        if hasattr(style, 'style_id'):
            try:
                style.style_id = style_id
            except: pass
        return style
    except KeyError:
        pass




    try:

        builtin_name = style_id.replace("BodyText", "Body Text ").replace("FirstIndent", "First Indent").strip()

        if style_id == 'BodyTextIndent': builtin_name = 'Body Text Indent'
        elif style_id == 'BodyTextFirstIndent': builtin_name = 'Body Text First Indent'
        elif style_id == 'BodyText': builtin_name = 'Body Text'
        elif style_id == 'Normal': builtin_name = 'Normal'
        
        try:
            style = doc.styles[builtin_name]
            style.name = name_cn
            return style
        except KeyError:

            style = doc.styles.add_style(builtin_name, WD_STYLE_TYPE.PARAGRAPH)
            style.style_id = style_id
            style.name = name_cn
            if base_style:
                try:
                    style.base_style = doc.styles[base_style]
                except KeyError: pass
            return style
            
    except Exception as e:
        print(f"创建样式失败 {name_cn} ({style_id}): {e}")

        style = doc.styles.add_style(name_cn, WD_STYLE_TYPE.PARAGRAPH)
        return style


def generate_reference_docx(output_path: str) -> bool:
    

       
    try:

        doc = Document()
        


        normal = _create_style_with_id(doc, '正文', 'Normal')
        normal.font.name = '宋体'
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        normal.font.size = Pt(10.5)
        


        body_text = _create_style_with_id(doc, '正文文本', 'BodyText', 'Normal')
        body_text.font.name = '宋体'
        body_text._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        


        body_indent = _create_style_with_id(doc, '正文文本缩进', 'BodyTextIndent', 'Normal')
        body_indent.paragraph_format.left_indent = Cm(0.74)
        


        

        body_first_indent = _create_style_with_id(doc, '正文首行缩进', 'BodyTextFirstIndent', 'Normal')
        body_first_indent.paragraph_format.first_line_indent = Cm(0.74)
        

        for i in range(1, 10):
            heading_id = f'Heading{i}'
            heading_name = f'Heading {i}'
            
            try:
                heading = doc.styles[heading_name]
            except KeyError:
                heading = doc.styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)
            

            if hasattr(heading, 'style_id'):
                 heading.style_id = heading_id
            

            heading.font.name = '黑体'
            heading._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            heading.font.bold = True

            sizes = [22, 16, 14, 12, 11, 10.5, 10.5, 10.5, 10.5]
            heading.font.size = Pt(sizes[i-1])
        


        caption = _create_style_with_id(doc, '题注', 'Caption', 'Normal')
        caption.font.size = Pt(9)
        caption.paragraph_format.alignment = 1
        

        doc.add_paragraph()
        

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"生成参考文档失�? {e}")
        return False


def get_or_create_reference_docx(resources_dir: str = None) -> str:
    







       
    if resources_dir is None:

        resources_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "resources"
        )
    
    ref_path = os.path.join(resources_dir, "default_reference.docx")
    

    if os.path.exists(ref_path):
        try:

            Document(ref_path)
            return ref_path
        except Exception:

            pass
    

    if generate_reference_docx(ref_path):
        return ref_path
    
    return None


if __name__ == "__main__":

    import sys
    

    script_dir = os.path.dirname(os.path.abspath(__file__))
    resources_dir = os.path.join(os.path.dirname(script_dir), "resources")
    output_path = os.path.join(resources_dir, "default_reference.docx")
    
    print(f"正在生成参考文�? {output_path}")
    if generate_reference_docx(output_path):
        print("�?参考文档生成成功！")
        

        doc = Document(output_path)
        styles = [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        print(f"  包含 {len(styles)} 个段落样�?)
        print(f"  包含样式: {', '.join(styles[:10])}...")
    else:
        print("�?参考文档生成失�?)
        sys.exit(1)
