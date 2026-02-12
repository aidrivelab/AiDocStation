# -*- coding: utf-8 -*-
"""
@File    : local_bridge/utils/docx_processor.py
@Desc    : AiDoc Station Lite 核心模块 - 赋能高效文档协作与智能排版处�?
@Author  : AIDriveLab Team
@Create  : 2026-02-09 21:12:43
@Version : V0.2.6
@Copyright: ©AIDriveLab Inc. All Rights Reserved.
"""

import os
import io
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..utils.logging import log
from ..constants.word_styles import apply_style_safe, get_english_style_name


class DocxProcessor:
                                          
    
    @staticmethod
    def normalize_first_paragraph_style(
        docx_bytes: bytes,
        target_style: str = "Body Text"
    ) -> bytes:
        








           
        try:

            doc = Document(io.BytesIO(docx_bytes))
            

            modified_count = 0
            

            for paragraph in doc.paragraphs:

                if paragraph.style and paragraph.style.name == "First Paragraph":

                    paragraph.style = target_style
                    modified_count += 1
                    log(f"Changed paragraph style from 'First Paragraph' to '{target_style}'")
            

            if modified_count > 0:
                log(f"Total {modified_count} paragraph(s) changed from 'First Paragraph' to '{target_style}'")
            else:
                log("No 'First Paragraph' style found in document")
            

            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            
            return output_stream.read()
            
        except Exception as e:
            log(f"Failed to process DOCX styles: {type(e).__name__}: {e}")

            return docx_bytes

    @staticmethod
    def add_table_borders(doc: Document) -> int:
        







           
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            tbl = table._tbl
            

            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            

            tblBorders = OxmlElement('w:tblBorders')
            

            border_attrs = {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '0',
                qn('w:color'): '000000'
            }
            

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                for attr, value in border_attrs.items():
                    border.set(attr, value)
                tblBorders.append(border)
            

            old_borders = tblPr.find(qn('w:tblBorders'))
            if old_borders is not None:
                tblPr.remove(old_borders)
            

            tblPr.append(tblBorders)
        
        return table_count

    @staticmethod
    def set_table_autofit(doc: Document) -> int:
        












           
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            tbl = table._tbl
            

            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            

            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            
            old_tblW = tblPr.find(qn('w:tblW'))
            if old_tblW is not None:
                tblPr.remove(old_tblW)
            tblPr.insert(0, tblW)
            

            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'autofit')
            
            old_layout = tblPr.find(qn('w:tblLayout'))
            if old_layout is not None:
                tblPr.remove(old_layout)
            tblPr.append(tblLayout)
            

            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    

                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), '0')
                    tcW.set(qn('w:type'), 'auto')
                    
                    old_tcW = tcPr.find(qn('w:tcW'))
                    if old_tcW is not None:
                        tcPr.remove(old_tcW)
                    tcPr.insert(0, tcW)
        
        return table_count

    @staticmethod
    def apply_table_text_style(doc: Document, target_style: str) -> int:
        








           
        if not target_style:
            return 0
        

        english_style = get_english_style_name(target_style)
        
        table_count = 0
        success_count = 0
        fail_count = 0
        
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:

                        DocxProcessor._ensure_style_exists(doc, target_style)
                        

                        if paragraph.style.name != english_style and paragraph.style.name != target_style:
                            if apply_style_safe(paragraph, target_style):
                                success_count += 1
                            else:
                                fail_count += 1
        
        if fail_count > 0:
            log(f"[Table Style] Failed to apply '{target_style}' to {fail_count} paragraph(s), succeeded: {success_count}")
        elif success_count > 0:
            log(f"[Table Style] Applied '{target_style}' to {success_count} paragraph(s)")
        
        return table_count

    @staticmethod
    def apply_body_text_style(doc: Document, target_style: str) -> int:
        








           
        if not target_style:
            return 0
            
        modified_count = 0
        


        source_styles = [
            "Normal", "Body Text", "Body Text 2", "Body Text 3",
            "Body Text Indent", "Body Text Indent 2", "Body Text Indent 3",
            "Body Text First Indent", "Body Text First Indent 2",
            "First Paragraph", "Compact", "Standard",

            "正文", "正文文本", "正文文本2", "正文文本3",
            "正文文本缩进", "正文文本缩进2", "正文文本缩进3",
            "正文首行缩进", "正文首行缩进2",
        ]
        

        target_english = get_english_style_name(target_style)


        for paragraph in doc.paragraphs:

            if paragraph.style and (paragraph.style.name.startswith("Heading") or "标题" in paragraph.style.name):
                continue

            if paragraph._p.getparent().tag.endswith('tc'):
                continue

            if DocxProcessor._contains_image(paragraph):
                continue
                
            if paragraph.style:
                s_name = paragraph.style.name
                

                if s_name in source_styles and s_name != target_style and s_name != target_english:

                    DocxProcessor._ensure_style_exists(doc, target_style)
                    

                    if apply_style_safe(paragraph, target_style):

                        DocxProcessor._strip_direct_indent(paragraph)
                        modified_count += 1
        
        return modified_count

    @staticmethod
    def _strip_direct_indent(paragraph) -> None:
        


           
        try:
            from docx.oxml.ns import qn
            pPr = paragraph._p.find(qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    pPr.remove(ind)
                    log("[V5.5] Stripped direct w:ind from paragraph.")
        except Exception as e:
            log(f"[V5.5] Failed to strip w:ind: {e}")

    @staticmethod
    def apply_custom_processing(
        docx_bytes: bytes,
        disable_first_para_indent: bool = False,
        target_style: str = "Body Text",
        add_table_borders: bool = True,
        table_autofit: bool = True,
        table_text_style: str = "正文",
        body_style: str = "正文",
        image_style: str = "正文",
        image_scale: int = 100,
        table_line_height_rule: str = "1.0"
    ) -> bytes:
        
















           
        try:

            doc = Document(io.BytesIO(docx_bytes))
            modified = False
            

            if body_style:
                count = DocxProcessor.apply_body_text_style(doc, body_style)
                if count > 0:
                    log(f"Applied body style '{body_style}' to {count} paragraph(s)")
                    modified = True


            if disable_first_para_indent and not body_style:
                modified_count = 0
                for paragraph in doc.paragraphs:
                    if paragraph.style and paragraph.style.name == "First Paragraph":
                        paragraph.style = target_style
                        modified_count += 1
                
                if modified_count > 0:
                    log(f"Total {modified_count} paragraph(s) changed from 'First Paragraph' to '{target_style}'")
                    modified = True
            

            if table_text_style:
                count = DocxProcessor.apply_table_text_style(doc, table_text_style)
                if count > 0:
                    log(f"Applied style '{table_text_style}' to {count} table(s)")
                    modified = True


            if add_table_borders:
                table_count = DocxProcessor.add_table_borders(doc)
                if table_count > 0:
                    log(f"Added borders to {table_count} table(s)")
                    modified = True
            

            if table_autofit or table_line_height_rule:
                count = 0
                if table_autofit:
                    count = DocxProcessor.set_table_autofit(doc)
                
                if table_line_height_rule:
                    DocxProcessor.apply_table_line_spacing(doc, table_line_height_rule)
                    
                if count > 0:
                    log(f"Applied table formatting (autofit={table_autofit}, spacing={table_line_height_rule})")
                    modified = True


            if image_style:
                count = DocxProcessor.apply_image_style(doc, image_style)
                if count > 0:
                    log(f"Applied image style '{image_style}' to {count} paragraph(s)")
                    modified = True
            

            if image_scale:
                try:



                    limit_w = 5760000
                    limit_h = 7200000
                    
                    try:
                        section = doc.sections[0]


                        page_w = section.page_width
                        page_h = section.page_height
                        

                        left_m = section.left_margin
                        right_m = section.right_margin
                        top_m = section.top_margin
                        bottom_m = section.bottom_margin
                        


                        safety_ratio = 0.99
                        if page_w and left_m is not None and right_m is not None:
                             limit_w = int((page_w - left_m - right_m) * safety_ratio)
                        
                        if page_h and top_m is not None and bottom_m is not None:
                             limit_h = int((page_h - top_m - bottom_m) * safety_ratio)
                             
                        log(f"[ImageScale] Dynamic Page Limit: W={limit_w} EMU, H={limit_h} EMU")
                    except Exception as dim_err:
                        log(f"[ImageScale] Failed to detect page dimensions, using defaults: {dim_err}")

                    from lxml import etree
                    

                    nsmap = {
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                    }
                    
                    scale_count = 0
                    user_scale_factor = image_scale / 100.0
                    

                    for paragraph in doc.paragraphs:

                        extents = paragraph._p.findall('.//wp:extent', namespaces=nsmap)
                        for ext in extents:
                            cx_str = ext.get('cx')
                            cy_str = ext.get('cy')
                            
                            if cx_str and cy_str:
                                try:
                                    orig_cx = int(cx_str)
                                    orig_cy = int(cy_str)
                                    

                                    new_cx = int(orig_cx * user_scale_factor)
                                    new_cy = int(orig_cy * user_scale_factor)
                                    


                                    if new_cx > limit_w:
                                        ratio_w = limit_w / new_cx
                                        new_cx = int(new_cx * ratio_w)
                                        new_cy = int(new_cy * ratio_w)
                                    

                                    if new_cy > limit_h:
                                        ratio_h = limit_h / new_cy
                                        new_cx = int(new_cx * ratio_h)
                                        new_cy = int(new_cy * ratio_h)
                                    

                                    if new_cx != orig_cx or new_cy != orig_cy:
                                        ext.set('cx', str(new_cx))
                                        ext.set('cy', str(new_cy))
                                        scale_count += 1
                                        




                                        pass 

                                except Exception:
                                    pass
                        




                        a_extents = paragraph._p.findall('.//a:ext', namespaces=nsmap)
                        for a_ext in a_extents:
                             cx_str = a_ext.get('cx')
                             cy_str = a_ext.get('cy')
                             if cx_str and cy_str:
                                try:
                                    orig_cx = int(cx_str)
                                    orig_cy = int(cy_str)
                                    

                                    new_cx = int(orig_cx * user_scale_factor)
                                    new_cy = int(orig_cy * user_scale_factor)
                                    
                                    if new_cx > limit_w:
                                        ratio_w = limit_w / new_cx
                                        new_cx = int(new_cx * ratio_w)
                                        new_cy = int(new_cy * ratio_w)
                                    if new_cy > limit_h:
                                        ratio_h = limit_h / new_cy
                                        new_cx = int(new_cx * ratio_h)
                                        new_cy = int(new_cy * ratio_h)
                                        
                                    if new_cx != orig_cx or new_cy != orig_cy:
                                        a_ext.set('cx', str(new_cx))
                                        a_ext.set('cy', str(new_cy))
                                except: pass
                    
                    if scale_count > 0:
                        log(f"Scaled {scale_count} image(s) with limits W<{limit_w} H<{limit_h}")
                        modified = True
                except Exception as scale_err:
                    log(f"Failed to scale images: {scale_err}")
            
            if modified:

                output_stream = io.BytesIO()
                doc.save(output_stream)
                output_stream.seek(0)
                return output_stream.read()
            
            return docx_bytes
            
        except Exception as e:
            log(f"Failed to apply custom processing: {type(e).__name__}: {e}")
            return docx_bytes

    @staticmethod
    def _contains_image(paragraph) -> bool:
                          
        p_xml = paragraph._p.xml

        return '<w:drawing' in p_xml or '<w:pict' in p_xml

    @staticmethod
    def apply_table_line_spacing(doc: Document, line_height_rule: str) -> None:
        





           
        from docx.enum.text import WD_LINE_SPACING
        
        spacing_map = {
            "1.0": WD_LINE_SPACING.SINGLE,
            "1.5": WD_LINE_SPACING.ONE_POINT_FIVE,
            "2.0": WD_LINE_SPACING.DOUBLE
        }
        rule = spacing_map.get(line_height_rule, WD_LINE_SPACING.SINGLE)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.line_spacing_rule = rule
    @staticmethod
    def _ensure_style_exists(doc: Document, style_name: str) -> None:
        


           
        english_name = get_english_style_name(style_name)
        

        try:
            if doc.styles[style_name]: return
        except KeyError:
            try:
                if doc.styles[english_name]: return
            except KeyError:
                pass




        current_dir = os.path.dirname(os.path.abspath(__file__))
        bridge_root = os.path.dirname(os.path.dirname(current_dir))
        ref_path = os.path.join(bridge_root, "local_bridge", "pandoc", "Reference-document.docx")
        
        if not os.path.exists(ref_path):
            log(f"[DocxProcessor] Reference document not found at {ref_path}, fallback to shell.")

            try:
                from docx.enum.style import WD_STYLE_TYPE
                new_style = doc.styles.add_style(english_name, WD_STYLE_TYPE.PARAGRAPH)
                new_style.base_style = doc.styles['Normal']
            except: pass
            return

        try:

            ref_doc = Document(ref_path)
            ref_styles = ref_doc.styles
            
            target_ref_style = None
            try:
                target_ref_style = ref_styles[style_name]
            except KeyError:
                try:
                    target_ref_style = ref_styles[english_name]
                except KeyError:
                    pass
            
            if target_ref_style:


                from lxml import etree
                style_xml = target_ref_style._element

                new_style_element = etree.fromstring(etree.tostring(style_xml))
                doc.styles._element.append(new_style_element)
                log(f"[DocxProcessor] Restored style '{style_name}' from reference document container.")
            else:
                log(f"[DocxProcessor] Style '{style_name}' not found even in reference doc! Fallback to Normal shell.")
                from docx.enum.style import WD_STYLE_TYPE
                new_style = doc.styles.add_style(english_name, WD_STYLE_TYPE.PARAGRAPH)
                new_style.base_style = doc.styles['Normal']
        except Exception as e:
            log(f"[DocxProcessor] Critical error during style restoration: {e}")
    @staticmethod
    def apply_image_style(doc: Document, target_style: str) -> int:
        

           
        if not target_style:
            return 0
            
        modified_count = 0
        for paragraph in doc.paragraphs:
            if DocxProcessor._contains_image(paragraph):

                DocxProcessor._ensure_style_exists(doc, target_style)

                if apply_style_safe(paragraph, target_style):

                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    modified_count += 1
        return modified_count
